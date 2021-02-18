#! python3

# Custom Invitations as Word Documents.py -  a program that would generate
# a Word document with custom invitations from a text file of guest names

import docx

doc = docx.Document()

# open the file 'guests.txt' and read all line of the text file
file_object = open('guests.txt', 'r')
names = file_object.readlines()

# each line of the text file is name of the guest so, add them to invitation format 'docx'
for name in names:
	doc.add_paragraph('It would be a pleasure to have the company of')
	doc.add_paragraph(name[:-1])
	doc.add_paragraph('at 11010 Memory Lane on the evening of')
	doc.add_paragraph('April 1st')
	doc.add_paragraph('at 7 o\'clock')

# search for last line of each invitation and add break line
for paragraph_num in range(4, len(doc.paragraphs), 5):
	doc.paragraphs[paragraph_num].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

# save the document as 'invitation.docx'
doc.save('invitation.docx')